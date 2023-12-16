# Copyright (c) shiqing. All rights reserved.
# !/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Author   : nan.qin
# @Company  : SHIQING TECH
# @Time     : 2022/09/01 14:28
# @File     : error_analysis.py

import os
import shutil
import cv2
import copy
from turtle import width
import numpy as np
import sys
from scipy.stats import entropy
sys.path.append(os.getcwd())
import tqdm
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import matplotlib.pyplot as plt

try:
    from utils import Object, process_im, Similarity
    from utils import extract_scale
    from simulator import Simulation
    from simulator import error_factory  # , CosineSimiarity, L1Simiarity, L2Simiarity
except:
    from onnx_converter.utils import Object, process_im, Similarity # type: ignore
    from onnx_converter.utils import extract_scale # type: ignore
    from onnx_converter.simulator import Simulation # type: ignore
    from onnx_converter.simulator import error_factory # type: ignore


scale_lambda = lambda data, scale: data * scale.reshape(1, -1, 1, 1) if isinstance(scale, np.ndarray) else data * scale


class ErrorThreshold():
    def __init__(self):
        self.cos_sim_decrease_threshold_1 = 0.05  # Single layer cosine similarity drop threshold
        self.cos_sim_decrease_threshold_2 = 0.01  # Single layer cosine similarity drop threshold with layer inputs\
                                                  # been set as the float inference results
        self.clip_ratio_threshold = 0.01          # layer's output clip ratio threshold
        self.weight_noise_ratio_threshold = 0.02  #
        self.quan_weight_abs_mean_threshold = 10  #

class ErrorAnalyzer(Simulation): # type: ignore
    def __init__(self, **kwargs):
        super(ErrorAnalyzer, self).__init__(**kwargs)
        # self.dataset = kwargs['dataset']
        # self.quan_graph = kwargs['quan_graph']
        # self.onnx_infer = kwargs['onnx_infer']
        # self.set_graph(kwargs['quan_graph'])
        # self.weight_distribution_dict = self.calc_weight_distribution()
        # self.featuremap_distribution_dict = self.init_featuremap_info_dict()
        # self.layer_error_info = self.init_layer_error_list()
        # self.layer_name_idx_dict = self.create_layer_name_idx_dict()
        self.weight_distribution_dict = []
        self.featuremap_distribution_dict = []
        self.layer_error_info = []
        self.layer_name_idx_dict = []
        # self.create_output_folder()
        # self.output_dir = kwargs['output_dir']
        self.output_dir = "work_dir/error_analyzer_test"
        
        self.error_threshold = ErrorThreshold()
        
        self.n_samples = 0

    def create_layer_name_idx_dict(self):
        layer_name_dict = dict()
        for i, layer in enumerate(self.get_layers()):
            layer_name_dict.update({layer.get_layer_name():i})
        return layer_name_dict

    def get_layer_idx_by_name(self, layer_name):
        assert layer_name in self.layer_name_idx_dict.keys(), \
            "Error, layer name not in quan_graph" # type: ignore
        if layer_name in self.layer_name_idx_dict.keys(): # type: ignore
            return self.layer_name_idx_dict[layer_name]
        else:
            return None

    def get_layer_process_scale(self, layer_idx):
        process_scale = self.get_layers()[layer_idx].get_ops_setting()['setting']['process_scale']
        return process_scale

    # def set_layer_process_scale(self, quan_graph, layer_idx, process_scale):
    #     setting = quan_graph.get_layers()[layer_idx].get_ops_setting()
    #     setting['setting']['process_scale'] = process_scale
    #     quan_graph.get_layers()[layer_idx].set_ops_setting(setting)

    def calc_clip_ratio(self,data, clip_range):
        size = data.size
        clip_size = np.sum(data <= clip_range[0])
        clip_size += np.sum(data >= clip_range[1])
        return clip_size / size

    def create_output_folder(self):
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        else:
            shutil.rmtree(self.output_dir)
            os.makedirs(self.output_dir)            
        # weight_folder = os.path.join(self.output_dir, 'weights')
        # featuremap_folder = os.path.join(self.output_dir, 'featuremaps')
        layer_folder = os.path.join(self.output_dir, 'layers')
        # if not os.path.exists(weight_folder):
        #     os.mkdir(weight_folder)
        # if not os.path.exists(featuremap_folder):
        #     os.mkdir(featuremap_folder)
        if not os.path.exists(layer_folder):
            os.mkdir(layer_folder)
        self.log_file_name = os.path.join(self.output_dir, 'model_error_analysis.docx')

    def init_featuremap_info_dict(self):
        self.featuremap_histogram_bins = 255
        featuremap_info_dict = dict()
        for layer in self.get_layers():
            bit_select = layer.get_ops_setting()['setting']['bit_select']
            bits_dict = layer.get_ops_setting()['setting']['bits_dict']
            featuremap_histogram_bins = 8 #bits_dict[bit_select].__name__.split('int')[-1]
            featuremap_histogram_bins = 2**int(featuremap_histogram_bins) - 1
            # featuremap_info_dict["featuremap_histogram_bins"] = featuremap_histogram_bins
            layer_input_names = layer.get_onnx_input_name()
            layer_output_names = layer.get_onnx_output_name()
            for name in layer_input_names:
                if not name in featuremap_info_dict.keys():
                    featuremap_info_dict.update({name:dict(featuremap_histogram_bins=featuremap_histogram_bins)})
            for name in layer_output_names:
                if not name in featuremap_info_dict.keys():
                    featuremap_info_dict.update({name:dict(featuremap_histogram_bins=featuremap_histogram_bins)})
        for key in featuremap_info_dict.keys():
            featuremap_histogram_bins = featuremap_info_dict[key]["featuremap_histogram_bins"]
            featuremap_info_dict[key].update({
                "mean_onnx":0,
                "std_onnx":0,
                "abs_mean_onnx":0,
                "min_onnx":1e15,
                "max_onnx":-1e15,
                "mean_quant":0,
                "std_quant":0,
                "abs_mean_quant":0,
                "min_quant":1e15,
                "max_quant":-1e15,
                "abs_mean_noise":0,
                "mean_noise":0,
                "clip_ratio": 0,
                "txme_clip_ratio": 0,
                "txme_lr_range": [0, 0, 0],
                "std_noise":0,
                "min_noise":1e15,
                "max_noise":-1e15,
                "cosine_error":0,
                "L1_error":0,
                "L2_error":0,
                "L1_error_no_input_error":0,
                "L2_error_no_input_error":0,
                "cosine_error_no_input_error":0,
                "entropy": 0,
                # "L1_error_no_input_error_floatscale":0,
                # "L2_error_no_input_error_floatscale":0,
                # "cosine_error_no_input_error_floatscale":0,
                "histogram_onnx":[0]*featuremap_histogram_bins,
                "histogram_quant":[0]*featuremap_histogram_bins,
                "histogram_noise":[0]*featuremap_histogram_bins,
                "clip_calculated":False,
                "count":0})
        return featuremap_info_dict

    def init_layer_error_list(self):
        layer_error_info = list()
        for layer in self.get_layers():
            layer_name = layer.get_layer_name()
            layer_type = layer.get_layer_type()
            layer_idx = layer.get_idx()
            # if layer_type == 'data':
            #     continue
            layer_error_info.append({'layer_idx': layer_idx, 'layer_name':layer_name, "layer_type":layer_type, 'weights':dict(), 'inputs':list(), 'outputs':list()})
            layer_input_names = layer.get_onnx_input_name()
            layer_output_names = layer.get_onnx_output_name()
            for name in layer_input_names:
                layer_error_info[-1]['inputs'].append({'name':name})
            for name in layer_output_names:
                layer_error_info[-1]['outputs'].append({'name':name})
        return layer_error_info

    def calc_weight_distribution(self):
        self.featuremap_histogram_bins=255
        weight_distribution_dict, w_smaller_dict = dict(), dict()
        for layer in self.get_layers():
            if not layer.layer_type in ["conv", "convtranspose", "gemm", "fc", "depthwiseconv", "matmul"]:
                continue
            layer_name = layer.get_layer_name()
            bit_select = layer.get_ops_setting()['setting']['w_bit_select']
            bits_dict = layer.get_ops_setting()['setting']['bits_dict']
            featuremap_histogram_bins = 8 #bits_dict[bit_select].__name__.split('int')[-1]
            featuremap_histogram_bins = 2**int(featuremap_histogram_bins) - 1
            weight_f = layer.get_weight()
            weight_q = layer.get_qweight()
            weight_min, weight_max = np.min(weight_f), np.max(weight_f)
            weight_mean = np.mean(weight_f)
            qweight_abs_mean = np.mean(np.abs(weight_q))
            weight_std = np.std(weight_f)
            sw, zp = layer.get_w_scale()['scale'], layer.get_w_scale()['zero_point']
            oc = weight_f.shape[0]
            weight_f = weight_f.reshape(oc,-1)
            weight_q = weight_q.reshape(oc,-1)
            weight_perchn_L1 = np.sum(np.abs(weight_f), axis = 1)
            if isinstance(zp, np.ndarray):
                weight_requant = (weight_q - zp.reshape(weight_q.shape[0],1)) * sw.reshape(weight_q.shape[0],1)
            else:
                weight_requant = (weight_q - zp) * sw
            weight_requant_min, weight_requant_max = np.min(weight_requant), np.max(weight_requant)
            weight_L1_error = np.mean(np.abs(weight_f - weight_requant))/np.mean(np.abs(weight_f))
            weight_L1_error_perchannel = np.mean(np.abs(weight_f - weight_requant), axis=1)/np.mean(np.abs(weight_f), axis=1)
            abs_max = np.maximum(np.abs(weight_min), np.abs(weight_max))
            hist_fp, _ = np.histogram(weight_f, featuremap_histogram_bins, range=(-abs_max, abs_max))
            hist_quant, _ = np.histogram(weight_requant, featuremap_histogram_bins, range=(-abs_max, abs_max))
            hist_fp = hist_fp / np.sum(hist_fp)
            hist_quant = hist_quant / np.sum(hist_quant)

            noise =  weight_requant - weight_f
            noise_mean = np.mean(noise)
            noise_std = np.std(noise)
            noise_min = np.min(noise)
            noise_max = np.max(noise)
            noise_ratio = noise_std / weight_std
            abs_max_val = np.maximum(np.abs(noise_min), np.abs(noise_max))
            hist_noise, _ = np.histogram(noise, featuremap_histogram_bins, range=(-abs_max_val, abs_max_val))
            hist_noise = hist_noise/np.sum(hist_noise)
            
            weight_f_cpy = copy.deepcopy(weight_f).reshape(-1)
            w_smaller_1_2_s = weight_f_cpy[weight_f_cpy < 0.5 * layer.get_w_scale()['scale']].shape[0] / weight_f_cpy.shape[0]
            w_smaller_dict[layer_name] = dict(w_smaller_1_2_s=w_smaller_1_2_s)
            
            weight_distribution_dict.update({layer_name :
            {
                "histogram_fp":hist_fp,
                "histogram_quant": hist_quant,
                "weight_mean":weight_mean,
                "abs_max" : abs_max,
                "weight_std":weight_std,
                "weight_min":weight_min,
                "weight_max":weight_max,
                "weight_requant_min":weight_requant_min,
                "weight_requant_max":weight_requant_max,                
                "noise_mean":noise_mean,
                "noise_std":noise_std,
                "noise_ratio":noise_ratio,
                "noise_min":noise_min,
                "noise_max":noise_max,
                "weight_perchn_L1":weight_perchn_L1,
                "histogram_noise":hist_noise,
                "qweight_abs_mean":qweight_abs_mean,
                "weight_L1_error":weight_L1_error,
                "weight_L1_error_perchannel": weight_L1_error_perchannel,
                "w_0.5s": w_smaller_1_2_s,
            }})
            
        # import json
        # with open("work_dir/w_0.5s.json", "w") as json_file:
        #     json.dump(w_smaller_dict, json_file)
                
        return weight_distribution_dict
    
    def forward(self, in_data, acc_error=True, true_outputs=None, isp_data=False):
        
        layers_info_fp_input = dict(layer_names=[], layer_in_idx=[])
        layers = self.get_layers()
        layers_outs = list()
        self.n_samples += 1
        for idx, layer in enumerate(layers):
            #print(idx)
            layer_type = layer.get_layer_type()
            if layer_type == 'data':
                if isinstance(in_data, dict):
                    layer.forward(in_data[layer.get_onnx_output_name()[0]])
                else:
                    first_conv_layer = layers[layer.get_output_idx()[0]]  
                    if first_conv_layer.get_layer_type() in ["conv", "convtranspose", "depthwiseconv"]:
                        pad_t, pad_l, pad_b, pad_r = first_conv_layer.get_ops_setting()["attrs"][0]["pads"] 
                        if not isp_data:
                            N, C, H, W = in_data.shape
                            H += (pad_t + pad_b)
                            in_data_tmp = np.zeros([N, C, H, W], dtype=np.float32)   
                            in_data_tmp[:, :, pad_t:H-pad_b, :] = in_data
                            in_data = in_data_tmp                    
                    layer.forward(in_data)
            else:
                #onnx_node = layer.get_input_nodes()
                input_names = layer.get_onnx_input_name()
                output_names = layer.get_onnx_output_name()
                onnx_inputs = [true_outputs[x] for x in input_names] # type: ignore
                onnx_outputs = [true_outputs[x] for x in output_names] # type: ignore

                # acc-error mode
                if not acc_error:
                    layers_info_fp_input = dict(layer_names=[], layer_in_idx=[])
                    inputs, info = self.get_qtrue_qdata(idx, true_outputs, acc_error)
                else:
                    layer_name = layer.get_layer_name()
                    if layer_name in layers_info_fp_input["layer_names"]:
                        idx_ = layers_info_fp_input["layer_names"].index(layer_name) 
                        valid_idx = layers_info_fp_input["layer_in_idx"][idx_]  
                        inputs, info = self.get_qtrue_qdata(idx, true_outputs, acc_error, valid_idx=valid_idx) 
                    else:                
                        inputs, info = self.get_layer_input(idx)
                    
                # one layer inference 
                layer.forward(inputs)
                
                if layer.get_first_conv():
                    continue
                                
                ops_setting= layer.get_ops_setting()['setting']
                bitwidth = ops_setting['int_scale']
                process_scale = ops_setting['process_scale']
                method = ops_setting['method']
                if ops_setting["txme_saturation"] <= 0:
                    if layer_type in ["conv", "fc", "matmul", "gemm"]:
                        bitwidth = ops_setting['int_scale'] + 2

                process_scale = ops_setting['process_scale']
                method = ops_setting['method']

                outputs = layer.get_out_data()
                if layer_type not in ["split", "shuffle", "lstm", "gru"]:
                    outputs = [outputs[-1]] if isinstance(outputs, list) else [outputs]
                if layer_type in ["lstm", "gru"]:
                    outputs = outputs[:len(layer.get_scale())]

                quan_inputs = [x['output'] for x in inputs]
                quan_outputs = [x['output'] for x in outputs]
                input_scales, output_scales = list(), list()
                for i, x in enumerate(inputs):
                    # if process_scale=='intscale' and method.endswith('floatsymquan'):
                    #     input_scales.append({'outshift':x['out_shift'], 'outscale':x['out_scale'], "scale": layer.in_scale[i]['scale'], 'zero_point':layer.in_scale[i]['zero_point']})
                    # elif process_scale in ['rshiftscale', 'rrshiftscale']:
                    #     input_scales.append({"scale": layer.in_scale[i]['scale'], 'zero_point':layer.in_scale[i]['zero_point']})
                    # else:
                    #     input_scales.append({"scale": layer.in_scale[i]['scale'], 'zero_point':layer.in_scale[i]['zero_point']})
                    input_scales.append({"scale": layer.get_in_scale()[i]['scale'], 'zero_point':layer.get_in_scale()[i]['zero_point']})
                for i, x in enumerate(outputs):
                    if process_scale == 'intscale':
                        output_scales.append({'outshift':x['out_shift'], 'outscale':x['out_scale'], "scale": layer.get_scale()[i]['scale'], 'zero_point':layer.get_scale()[i]['zero_point']})
                    elif process_scale in ['rshiftscale', 'rrshiftscale']:# scale = layer.scales[i]['fscale']
                        output_scales.append({"scale": layer.get_scales()[i]['fscale'], 'zero_point':0})
                    elif process_scale == "preintscale":
                        output_scales.append({'outshift':x['out_shift'], 'outscale':x['out_scale'], "scale": layer.get_scale()[i]['scale'], 'zero_point':layer.get_scale()[i]['zero_point']})
                    else:
                        output_scales.append({"scale": layer.get_scale()[i]['scale'], 'zero_point':layer.get_scale()[i]['zero_point']})

                clip_range = None # clip calc condition, int-scale and
                if process_scale in ['intscale', 'shiftfloatscale', 'rshiftscale', 'rrshiftscale', 
                                     'shiftfloatscaletable', 'shiftfloatscaletable2float']:
                    if ops_setting['bit_select'] % 2 == 0:
                        clip_range = [0, np.int(2**(bitwidth)-1)] # type: ignore
                    else:
                        clip_range = [np.int(-2**(bitwidth-1)), np.int(2**(bitwidth-1)-1)] # type: ignore
                    # clip_range = [np.int(-2**(bitwidth-1) * output_scales[0]['outscale'] / 2**bitwidth),\
                    #     np.int((2**(bitwidth-1)-1) * output_scales[0]['outscale']  / 2**bitwidth) ]

                bit_select = layer.get_ops_setting()['setting']['w_bit_select']
                bits_dict = layer.get_ops_setting()['setting']['bits_dict']
                featuremap_histogram_bins = 8 #bits_dict[bit_select].__name__.split('int')[-1]
                featuremap_histogram_bins = 2**int(featuremap_histogram_bins) - 1

                layers_outs.append({
                    "layer_name" : layer.get_layer_name(),
                    "layer_type" : layer.get_layer_type(),
                    "input_names": input_names,
                    "output_names" : output_names,
                    "input_quant" : quan_inputs,
                    "input_onnx": onnx_inputs,
                    "output_quant" : quan_outputs,
                    "output_onnx":onnx_outputs,
                    "input_scales":input_scales,
                    "output_scales":output_scales,
                    "clip_range": clip_range,
                    "featuremap_histogram_bins": featuremap_histogram_bins})
                # print(layer.get_layer_name())

        for i, layer_outputs in enumerate(layers_outs):
            layer_name = layer_outputs['layer_name']
            layer_idx = self.get_layer_idx_by_name(layer_name)              
            for j, name in enumerate(layer_outputs['input_names']):
                if name not in self.network_input_names:
                    continue    
                # get onnx and quantized featuremap, and quantize-noise
                info = self.featuremap_distribution_dict[name]
                # if info['count'] == n_samples:
                #     continue
                data_quant = layer_outputs['input_quant'][j]
                data_onnx = layer_outputs['input_onnx'][j].astype(np.float32)
                scale, zp = layer_outputs['input_scales'][j]['scale'], layer_outputs['input_scales'][j]['zero_point']
                if data_quant.dtype != np.float32:
                    data_requant = layers[layer_idx].get_in_quantize()[0].get_dequan_data(data_quant)
                else:
                    data_requant=data_quant
                quant_noise = data_requant - data_onnx

                info = self.featuremap_distribution_dict[name]
                info['mean_onnx'] += np.mean(data_onnx)
                info['std_onnx'] += np.mean(data_onnx ** 2)
                info['abs_mean_onnx'] += np.mean(np.abs(data_onnx))
                info['min_onnx'] = np.minimum(info['min_onnx'], np.min(data_onnx))
                info['max_onnx'] = np.maximum(info['max_onnx'], np.max(data_onnx))
                info['mean_quant'] += np.mean(data_requant)
                info['std_quant'] += np.mean(data_requant ** 2)
                info['abs_mean_quant'] += np.mean(np.abs(data_requant))
                info['min_quant'] = np.minimum(info['min_quant'], np.min(data_requant))
                info['max_quant'] = np.maximum(info['max_quant'], np.max(data_requant))
                info['cosine_error'] += error_factory.get('Cosine')()(data_onnx, data_requant)/100 # type: ignore
                info['L1_error'] += error_factory.get('L1')()(data_onnx, data_requant)/100 # type: ignore
                info['L2_error'] += error_factory.get('L2')()(data_onnx, data_requant)/100 # type: ignore
                info['abs_mean_noise'] += np.mean(np.abs(quant_noise))
                info['mean_noise'] += np.mean(quant_noise)
                info['std_noise'] += np.mean(quant_noise ** 2)
                info['min_noise'] = np.minimum(info['min_noise'], np.min(quant_noise))
                info['max_noise'] = np.maximum(info['max_noise'], np.max(quant_noise))
                if layer_outputs['clip_range']:
                    info['clip_calculated'] = True
                    info['clip_ratio'] += self.calc_clip_ratio(data_quant, layer_outputs['clip_range'])
                # info['count']+=1
            for j, name in enumerate(layer_outputs['output_names']):
                info = self.featuremap_distribution_dict[name]
                # if info['count'] == n_samples:
                #     continue
                # get onnx and quantized featuremap, and quantize-noise
                data_quant = layer_outputs['output_quant'][j]
                data_onnx = layer_outputs['output_onnx'][j].astype(np.float32)
                scale, zp = layer_outputs['output_scales'][j]['scale'], layer_outputs['output_scales'][j]['zero_point']
                if data_quant.dtype != np.float32:
                    data_requant = layers[layer_idx].get_quantize()["feat"][f"so{j}"].get_dequan_data(data_quant)
                else:
                    data_requant=data_quant
                quant_noise = data_requant - data_onnx

                info['mean_onnx'] += np.mean(data_onnx)
                info['std_onnx'] += np.mean(data_onnx ** 2)
                info['abs_mean_onnx'] += np.mean(np.abs(data_onnx))
                info['min_onnx'] = np.minimum(info['min_onnx'], np.min(data_onnx))
                info['max_onnx'] = np.maximum(info['max_onnx'], np.max(data_onnx))
                info['mean_quant'] += np.mean(data_requant)
                info['std_quant'] += np.mean(data_requant ** 2)
                info['abs_mean_quant'] += np.mean(np.abs(data_requant))
                info['min_quant'] = np.minimum(info['min_quant'], np.min(data_requant))
                info['max_quant'] = np.maximum(info['max_quant'], np.max(data_requant))
                info['cosine_error'] += error_factory.get('Cosine')()(data_onnx, data_requant)/100 # type: ignore
                info['L1_error'] += error_factory.get('L1')()(data_onnx, data_requant)/100 # type: ignore
                info['L2_error'] += error_factory.get('L2')()(data_onnx, data_requant)/100 # type: ignore
                info['abs_mean_noise'] += np.mean(np.abs(quant_noise))
                info['mean_noise'] += np.mean(quant_noise)
                info['std_noise'] += np.mean(quant_noise ** 2)
                info['min_noise'] = np.minimum(info['min_noise'], np.min(quant_noise))
                info['max_noise'] = np.maximum(info['max_noise'], np.max(quant_noise))
                # info['entropy'] += entropy(data_requant.reshape(-1), data_onnx.reshape(-1), base=2, axis=-1)
                if layer_outputs['clip_range']:
                    info['clip_calculated'] = True
                    info['clip_ratio'] += self.calc_clip_ratio(data_quant, layer_outputs['clip_range'])
                
        if True: #not acc_error:
            for i, layer_outputs in enumerate(layers_outs):
                layer_name = layer_outputs['layer_name']
                layer_idx = self.get_layer_idx_by_name(layer_name)
                process_scale = self.get_layer_process_scale(layer_idx)
                #print("###### TEST: process_scale=%s"%process_scale)
                for j, name in enumerate(layer_outputs['output_names']):
                    info = self.featuremap_distribution_dict[name]
                    # get onnx and quantized featuremap, and quantize-noise
                    data_quant = layer_outputs['output_quant'][j]
                    data_onnx = layer_outputs['output_onnx'][j].astype(np.float32)
                    scale, zp = layer_outputs['output_scales'][j]['scale'], layer_outputs['output_scales'][j]['zero_point']

                    if data_quant.dtype != np.float32:
                        if process_scale in ['rshiftscale', 'rrshiftscale']:
                            data_requant = layers[layer_idx].get_quantize()["feat"][f"so{j}"].get_dequan_data(data_quant)
                        else:
                            data_requant = (data_quant - zp) * scale # requantize to float
                    else:
                        data_requant=data_quant
                    quant_noise = data_requant - data_onnx
                    # save error-metrics
                    info['cosine_error_no_input_error'] += error_factory.get('Cosine')()(data_onnx, data_requant)/100 # type: ignore
                    info['L1_error_no_input_error'] += error_factory.get('L1')()(data_onnx, data_requant)/100 # type: ignore
                    info['L2_error_no_input_error'] += error_factory.get('L2')()(data_onnx, data_requant)/100 # type: ignore

        for i, layer_outputs in enumerate(layers_outs):
            layer_name = layer_outputs['layer_name']
            layer_idx = self.get_layer_idx_by_name(layer_name)            
            featuremap_histogram_bins = layer_outputs["featuremap_histogram_bins"]
            for j, name in enumerate(layer_outputs['input_names']):
                if name not in self.network_input_names:
                    continue                
                data_quant = layer_outputs['input_quant'][j]
                data_onnx = layer_outputs['input_onnx'][j].astype(np.float32)
                scale, zp = layer_outputs['input_scales'][j]['scale'], layer_outputs['input_scales'][j]['zero_point']
                if data_quant.dtype != np.float32:
                    data_requant = layers[layer_idx].get_in_quantize()[0].get_dequan_data(data_quant)
                else:
                    data_requant=data_quant
                quant_noise = data_requant - data_onnx
                info = self.featuremap_distribution_dict[name]
                abs_max_val = np.maximum(np.maximum(np.abs(info['max_onnx']), np.abs(info['min_onnx'])), \
                    np.maximum(np.abs(info['max_quant']), np.abs(info['min_quant'])))
                info['histogram_onnx'] += np.histogram(data_onnx, featuremap_histogram_bins, range=(-abs_max_val, abs_max_val))[0].astype(np.float32)
                info['histogram_quant'] += np.histogram(data_requant, featuremap_histogram_bins, range=(-abs_max_val, abs_max_val))[0].astype(np.float32)
                abs_max_val = np.maximum(np.abs(info['max_noise']), np.abs(info['min_noise']))
                info['histogram_noise'] += np.histogram(quant_noise, featuremap_histogram_bins, range=(-abs_max_val, abs_max_val))[0].astype(np.float32)
                # info['count']+=1
            for j, name in enumerate(layer_outputs['output_names']):
                featuremap_histogram_bins = layer_outputs["featuremap_histogram_bins"]
                info = self.featuremap_distribution_dict[name]
                # if info['count'] == n_samples:
                #     continue
                if layers[layer_idx].get_layer_type() in ["conv", "convtranspose", "depthwiseconv"]:
                    out_shift = -layers[layer_idx].get_scales()[-1]["out_shift"]
                    data_after_shift = layers[layer_idx].get_out_data()[1]["output"] >> out_shift
                    min_v, max_v = np.min(data_after_shift.reshape(-1)), np.max(data_after_shift.reshape(-1))
                    if min_v < clip_range[0]:
                        d_value = clip_range[0] - min_v
                        info['txme_lr_range'][2] += d_value
                        if d_value > info['txme_lr_range'][0]:
                            info['txme_lr_range'][0] = d_value
                    if max_v > clip_range[1]:
                        d_value = max_v - clip_range[1] 
                        info['txme_lr_range'][2] += d_value
                        if d_value > info['txme_lr_range'][1]:
                            info['txme_lr_range'][1] = d_value
                    info['txme_clip_ratio'] += self.calc_clip_ratio(data_after_shift, clip_range)
                    
                if layers[layer_idx].get_layer_type() in ["concat"]:
                    indatas = layers[layer_idx].get_in_data()
                    # outdatas = layers[layer_idx].get_out_data()
                    scales = layers[layer_idx].get_scales()
                    for indata, scale in zip(indatas, scales):
                        data_after_shift = (indata['output'].astype(np.int32) * scale['out_scale']) >> scale['int_scale']
                        info['txme_clip_ratio'] += self.calc_clip_ratio(data_after_shift, clip_range)
                    # print("test")
                
                data_quant = layer_outputs['output_quant'][j]
                data_onnx = layer_outputs['output_onnx'][j].astype(np.float32)
                scale, zp = layer_outputs['output_scales'][j]['scale'], layer_outputs['output_scales'][j]['zero_point']
                if data_quant.dtype != np.float32:
                    data_requant = layers[layer_idx].get_quantize()["feat"][f"so{j}"].get_dequan_data(data_quant)
                else:
                    data_requant=data_quant
                quant_noise = data_requant - data_onnx
                info = self.featuremap_distribution_dict[name]
                abs_max_val = np.maximum(np.maximum(np.abs(info['max_onnx']), np.abs(info['min_onnx'])), \
                    np.maximum(np.abs(info['max_quant']), np.abs(info['min_quant'])))
                if len(info['histogram_onnx']) != featuremap_histogram_bins:
                    featuremap_histogram_bins = len(info['histogram_onnx'])
                info['histogram_onnx'] += np.histogram(data_onnx, featuremap_histogram_bins, range=(-abs_max_val, abs_max_val))[0].astype(np.float32)
                info['histogram_quant'] += np.histogram(data_requant, featuremap_histogram_bins, range=(-abs_max_val, abs_max_val))[0].astype(np.float32)
                abs_max_val = np.maximum(np.abs(info['max_noise']), np.abs(info['min_noise']))
                info['histogram_noise'] += np.histogram(quant_noise, featuremap_histogram_bins, range=(-abs_max_val, abs_max_val))[0].astype(np.float32)
        
        # return dict(results=self.process_output(true_outputs))
        
    def calc_featuremap_distribution(self):
        n_samples = self.n_samples

        for key in self.featuremap_distribution_dict.keys(): # type: ignore
            info = self.featuremap_distribution_dict[key]
            info['mean_onnx'] /= n_samples
            info['std_onnx'] = np.sqrt(info['std_onnx'] / n_samples - info['mean_onnx']**2)
            info['abs_mean_onnx'] /= n_samples
            info['mean_quant'] /= n_samples
            info['std_quant'] = np.sqrt(info['std_quant'] / n_samples - info['mean_quant']**2)
            info['abs_mean_quant'] /= n_samples
            info['cosine_error'] /= n_samples
            info['L1_error'] /= n_samples
            info['L2_error'] /= n_samples
            info['abs_mean_noise'] /= n_samples
            info['mean_noise'] /= n_samples
            info['std_noise'] = np.sqrt(info['std_noise'] / n_samples - info['mean_noise']**2)
            info['noise_ratio'] = info['std_noise'] / info['std_onnx'] if info['std_onnx'] > 0 else 0
            info['entropy'] /= n_samples
            info['txme_clip_ratio'] /= n_samples
            info['txme_lr_range'][2] /= (2 * n_samples)
            if info["clip_calculated"] == True:
                info["clip_ratio"] /=n_samples
            # info['count'] = 0

        for key in self.featuremap_distribution_dict.keys(): # type: ignore
            info = self.featuremap_distribution_dict[key]
            if 'cosine_error_no_input_error' in info.keys():
                info['cosine_error_no_input_error'] /= n_samples
                info['L1_error_no_input_error'] /= n_samples
                info['L2_error_no_input_error'] /= n_samples
                # info['cosine_error_no_input_error_floatscale'] /= n_samples
                # info['L1_error_no_input_error_floatscale'] /= n_samples
                # info['L2_error_no_input_error_floatscale'] /= n_samples
        # self.logger.info("featuremap distribution info")
        # for key in self.featuremap_distribution_dict.keys():
        #     info = self.featuremap_distribution_dict[key]
        #     # log_info = "output node name is: {}, cosine error: {}, cosine_error_no_input_error: {}, clip_ratio: {}, \
        #     #     entropy is: {}".format(key, info['cosine_error'], info['cosine_error_no_input_error'], info['clip_ratio'], info['entropy'])
        #     log_info = "output node name is: {}, cosine error: {}, cosine_erro: {}, clip_ratio: {} \
        #         ".format(key, info['cosine_error'], info['cosine_error_no_input_error'], info['clip_ratio'])
        #     self.logger.info(log_info)

        
        for key in self.featuremap_distribution_dict.keys(): # type: ignore
            info=self.featuremap_distribution_dict[key]
            info['histogram_onnx'] /= np.sum(info['histogram_onnx'])
            info['histogram_quant'] /= np.sum(info['histogram_quant'])
            info['histogram_noise'] /= np.sum(info['histogram_noise'])

    def generate_layer_error_info(self):
        for i, info in enumerate(self.layer_error_info):
            layer_name = info['layer_name']
            layer_type = info['layer_type']
            if layer_name in self.weight_distribution_dict.keys(): # type: ignore
                info['weights'].update(self.weight_distribution_dict[layer_name])
            for input in info['inputs']:
                name = input['name']
                input.update(self.featuremap_distribution_dict[name])
            for output in info['outputs']:
                name = output['name']
                output.update(self.featuremap_distribution_dict[name])

    def save_result_figures(self):
        def draw_hist(x, hist_onnx, hist_quant, hist_onnx_range, hist_quant_range, hist_name, titles, bins=255):
            fig, (ax1, ax2) = plt.subplots(nrows=1, ncols=2, figsize=(16, 8)) # type: ignore
            ax1.plot(x, hist_onnx, color='blue', linestyle='-')
            ax1.set_ylim([0, np.max([hist_onnx.max(), hist_quant.max()])])
            ax1.axvline(x=0, linestyle='dashed', color='g')
            ax1.set_title(titles[0])
            ax2.plot(x, hist_quant, color='red', linestyle='-')
            ax2.set_ylim([0, np.max([hist_onnx.max(), hist_quant.max()])])
            ax2.axvline(x=0, linestyle='dashed', color='g')
            ax2.set_title(titles[1])
            ax1.annotate('min: {:.4f}\nmax: {:.4f}'.format(
                hist_onnx_range[0], hist_onnx_range[1]), xy=(0.7, 0.9), xycoords='axes fraction')
            ax2.annotate('min: {:.4f}\nmax: {:.4f}'.format(
                hist_quant_range[0], hist_quant_range[1]), xy=(0.7, 0.9), xycoords='axes fraction')                       
            plt.savefig(hist_name, dpi=300, bbox_inches='tight') 
            plt.close()
                    
        # draw per-layer info
        self.logger.info(" ### Save analysis images and log")
        layer_folder = os.path.join(self.output_dir, 'layers')

        for idx,info in tqdm.tqdm(enumerate(self.layer_error_info)):
            layer_name = info['layer_name'].replace('/', '_')
            layer_type = info['layer_type']
            # create folder
            current_dir = os.path.join(layer_folder, "{:04d}-{}".format(idx,layer_name))
            if not os.path.exists(current_dir):
                os.makedirs(current_dir)
            inputs_info = info['inputs']
            outputs_info = info['outputs']
            # draw input info
            for i,input_info in enumerate(inputs_info):
                featuremap_histogram_bins = input_info["featuremap_histogram_bins"]
                input_name = input_info['name'].replace('/', '_')
                hist_name = os.path.join(current_dir, "{}_input_{}-{}_histogram.png".format(layer_name,i,input_name))
                hist_onnx = input_info['histogram_onnx']
                hist_quant = input_info['histogram_quant']
                abs_max_val = np.maximum(np.maximum(np.abs(input_info['max_onnx']), np.abs(input_info['min_onnx'])), \
                    np.maximum(np.abs(input_info['max_quant']), np.abs(input_info['min_quant'])))
                x = np.linspace(-abs_max_val, abs_max_val, featuremap_histogram_bins)
                # plt.figure(figsize=(16, 8))
                # l1, = plt.plot(x, hist_onnx, 'r--')
                # l2, = plt.plot(x, hist_quant, 'b--')
                # plt.legend([l1, l2],['onnx', 'quantized'])
                # plt.savefig(hist_name, bbox_inches='tight')
                # plt.close()
                hist_onnx_range = [input_info['min_onnx'], input_info['max_onnx']]
                hist_quant_range = [input_info['min_quant'], input_info['max_quant']]
                if np.isnan(hist_onnx.max()) or np.isnan(hist_quant.max()):
                    continue                
                draw_hist(x, hist_onnx, hist_quant, hist_onnx_range, 
                          hist_quant_range, hist_name, 
                          titles=['onnx', 'quantized'], 
                          bins=featuremap_histogram_bins)
                # draw noise hist

                abs_max_val = np.maximum(np.abs(input_info['max_noise']), np.abs(input_info['min_noise']))
                x = np.linspace(-abs_max_val, abs_max_val, featuremap_histogram_bins)
                hist_noise = input_info['histogram_noise']
                plt.figure(figsize=(16, 8))
                plt.bar(x, np.sqrt(hist_noise), abs_max_val/512, color='blue')
                hist_name = os.path.join(current_dir, "{}_input_error_{}-{}_histogram-sqrt.png".format(layer_name,i,input_name))
                plt.savefig(hist_name, bbox_inches='tight')
                plt.close()
            # draw output info
            for i,output_info in enumerate(outputs_info):
                featuremap_histogram_bins = output_info["featuremap_histogram_bins"]
                input_name = output_info['name'].replace('/', '_')
                hist_name = os.path.join(current_dir, "{}_output_{}-{}_histogram.png".format(layer_name,i,input_name))
                hist_onnx = output_info['histogram_onnx']
                hist_quant = output_info['histogram_quant']
                abs_max_val = np.maximum(np.maximum(np.abs(output_info['max_onnx']), np.abs(output_info['min_onnx'])), \
                    np.maximum(np.abs(output_info['max_quant']), np.abs(output_info['min_quant'])))
                x = np.linspace(-abs_max_val, abs_max_val, featuremap_histogram_bins)
                # plt.figure(figsize=(16, 8))
                # l1, = plt.plot(x, hist_onnx, 'r--')
                # l2, = plt.plot(x, hist_quant, 'b--')
                # plt.legend([l1, l2],['onnx', 'quantized'])
                # plt.savefig(hist_name, bbox_inches='tight')
                # plt.close()
                hist_onnx_range = [output_info['min_onnx'], output_info['max_onnx']]
                hist_quant_range = [output_info['min_quant'], output_info['max_quant']]  
                if np.isnan(hist_onnx.max()) or np.isnan(hist_quant.max()):
                    continue          
                draw_hist(x, hist_onnx, hist_quant, hist_onnx_range,
                          hist_quant_range, hist_name, 
                          titles=['onnx', 'quantized'], 
                          bins=featuremap_histogram_bins)
                # draw noise hist

                abs_max_val = np.maximum(np.abs(output_info['max_noise']), np.abs(output_info['min_noise']))
                x = np.linspace(-abs_max_val, abs_max_val, featuremap_histogram_bins)
                hist_noise = output_info['histogram_noise']
                plt.figure(figsize=(16, 8))
                plt.bar(x, np.sqrt(hist_noise), abs_max_val/512, color='blue')
                hist_name = os.path.join(current_dir, "{}_output_error_{}-{}_histogram-sqrt.png".format(layer_name,i,input_name))
                plt.savefig(hist_name, bbox_inches='tight')
                plt.close()
            # draw weight_info (if conv-like)
            if not layer_type in ["conv", "convtranspose", "gemm", "fc", "depthwiseconv", "matmul"]:
                continue
            weight_info = info['weights']
            hist_fp, hist_quant = weight_info['histogram_fp'], weight_info['histogram_quant']
            hist_noise = weight_info["histogram_noise"]
            abs_max_val = np.maximum(np.abs(weight_info['weight_min']),np.abs(weight_info['weight_max']))
            hist_name = os.path.join(current_dir, "{}_weight_histogram.png".format(layer_name))
            x = np.linspace(-abs_max_val, abs_max_val, featuremap_histogram_bins) # type: ignore
            # plt.figure(figsize=(16, 8))
            # l1, = plt.plot(x, hist_fp, 'r--')
            # l2, = plt.plot(x, hist_quant, 'b--')
            # plt.legend([l1, l2],['fp', 'quantized'])
            # plt.savefig(hist_name, bbox_inches='tight')
            # plt.close()
            hist_fp_range = [weight_info['weight_min'], weight_info['weight_max']]
            hist_quant_range = [weight_info['weight_requant_min'], weight_info['weight_requant_max']]            
            if np.isnan(hist_fp.max()) or np.isnan(hist_quant.max()):
                continue            
            draw_hist(x, hist_fp, hist_quant, hist_fp_range,
                      hist_quant_range, hist_name, 
                      titles=['fp', 'quantized'], 
                      bins=featuremap_histogram_bins) # type: ignore
                        
            hist_name = os.path.join(current_dir, "{}_weight_noise_histogram.png".format(layer_name))
            abs_max_val = np.maximum(np.abs(weight_info['noise_min']), np.abs(weight_info['noise_max']))
            x = np.linspace(-abs_max_val, abs_max_val, featuremap_histogram_bins) # type: ignore
            plt.figure(figsize=(16, 8))
            plt.bar(x, hist_noise, abs_max_val/512, color='blue')
            plt.savefig(hist_name, bbox_inches='tight')
            plt.close()

            per_chn_L1 = weight_info["weight_perchn_L1"]
            img_name = os.path.join(current_dir, "{}_weight_per_channel_L1.png".format(layer_name))
            oc = per_chn_L1.size
            plt.figure(figsize=(16, 8))
            plt.bar(np.linspace(0,oc,oc), per_chn_L1, 0.1, color='blue')
            plt.savefig(img_name, bbox_inches='tight')
            plt.close()

            per_chn_L1_error = weight_info["weight_L1_error_perchannel"]
            img_name = os.path.join(current_dir, "{}_weight_L1_error_perchannel.png".format(layer_name))
            oc = per_chn_L1.size
            plt.figure(figsize=(16, 8))
            plt.bar(np.linspace(0,oc,oc), per_chn_L1_error, 0.1, color='blue')
            plt.savefig(img_name, bbox_inches='tight')
            plt.close()

        layer_name_list = [x['layer_name'] for x in self.layer_error_info]
        num_layers = len(layer_name_list)

        # save figures of error metrics
        x=range(0, num_layers, 1)
        plt.figure(figsize=(int(num_layers/8), 8))
        plt.bar(x, self.metrics_L1, 0.2, color='blue')
        plt.xticks(x, tuple(layer_name_list))
        plt.xticks(rotation=90, fontsize=8)
        plt.savefig(os.path.join(self.output_dir, "L1_error.png"), bbox_inches='tight')
        plt.close()

        plt.figure(figsize=(int(num_layers/8), 8))
        plt.bar(x, self.metrics_L2, 0.2, color='blue')
        plt.xticks(x, tuple(layer_name_list))
        plt.xticks(rotation=90, fontsize=8)
        plt.savefig(os.path.join(self.output_dir, "L2_error.png"), bbox_inches='tight')
        plt.close()

        plt.figure(figsize=(int(num_layers/8), 8))
        plt.bar(x, self.metrics_cosine, 0.2, color='blue')
        plt.xticks(x, tuple(layer_name_list))
        plt.xticks(rotation=90, fontsize=8)
        plt.savefig(os.path.join(self.output_dir, "Cosine_error.png"), bbox_inches='tight')
        plt.close()

        # plt.figure(figsize=(int(num_layers/8), 8))
        # plt.bar(x, self.metrics_noise_ratio, 0.2)
        # plt.xticks(x, tuple(layer_name_list))
        # plt.xticks(rotation=90, fontsize=8)
        # plt.savefig(os.path.join(self.output_dir, "noise_ratio_similarity.png"), bbox_inches='tight')
        # plt.close()

        plt.figure(figsize=(int(num_layers/8), 8))
        plt.bar(x, self.cosine_error_variance, 0.2, color='blue')
        plt.xticks(x, tuple(layer_name_list))
        plt.xticks(rotation=90, fontsize=8)
        plt.savefig(os.path.join(self.output_dir, "Cosine_error_variance.png"), bbox_inches='tight')
        plt.close()

    def calc_layer_error_metrics(self):
        # draw similarity metrics
        self.metrics_L1, self.metrics_L2, self.metrics_cosine, self.metrics_noise_ratio,self.metrics_cosine_acc_off = \
            list(), list(), list(), list(), list()
        for info in self.layer_error_info:
            output = info['outputs'][0]
            #layer_name_list.append(info['layer_name'])
            self.metrics_L1.append(output["L1_error"])
            self.metrics_L2.append(output["L2_error"])
            self.metrics_cosine.append(output['cosine_error'])
            self.metrics_cosine_acc_off.append(output['cosine_error_no_input_error'])
            self.metrics_noise_ratio.append(output["noise_ratio"])

        self.cosine_error_variance = [self.metrics_cosine[i]-self.metrics_cosine[i-1] for i in range(1,len(self.metrics_cosine))]
        self.cosine_error_variance.insert(0, 0)

        self.metrics_L1 = np.array(self.metrics_L1)
        self.metrics_L2 = np.array(self.metrics_L2)
        self.metrics_cosine = np.array(self.metrics_cosine)
        self.metrics_noise_ratio=np.array(self.metrics_noise_ratio)
        self.cosine_error_variance = np.array(self.cosine_error_variance)
        self.metrics_cosine_acc_off = np.array(self.metrics_cosine_acc_off)

    def low_accuracy_analysis(self):
        # find low similarity layers
        low_acc_layer_idx1 = set(np.argwhere(self.cosine_error_variance > self.error_threshold.cos_sim_decrease_threshold_1).flatten()) # type: ignore
        low_acc_layer_idx2 = set(np.argwhere(self.metrics_cosine_acc_off > self.error_threshold.cos_sim_decrease_threshold_2).flatten()) # type: ignore
        self.low_acc_layer_idx = list(low_acc_layer_idx1.union(low_acc_layer_idx2))
        self.low_acc_layer_idx.sort()
        # if len(self.low_acc_layer_idx):
        #     self.logger.info("The low accuracy layers are as below: ")
        #     for idx in self.low_acc_layer_idx:
        #         self.logger.info("{}, {}".format(idx, self.layer_error_info[idx]['outputs'][0]['cosine_error']))
        # analysis bad layers error
        self.low_acc_layers_error_source = list()
        for idx, layer_idx in enumerate(self.low_acc_layer_idx):
            self.low_acc_layers_error_source.append(dict({"Input Accumulated Error":0, "Weight Quantization Noise":0, "Data Clip":0}))
            layer_info = self.layer_error_info[layer_idx]
            layer_name = layer_info['layer_name']
            layer_type = layer_info['layer_type']
            layer_inputs_info = layer_info['inputs']
            layer_outputs_info = layer_info['outputs']
            weight_info = None
            if layer_type in ["conv", "convtranspose", "gemm", "fc", "depthwiseconv", "matmul"]:
                weight_info = layer_info['weights']
                qweight_abs_mean = weight_info['qweight_abs_mean']
                wegiht_noise_ratio = weight_info['noise_ratio']
            output_cos_error_acc_on = np.mean([x['cosine_error'] for x in layer_outputs_info])
            output_cos_error_acc_off = np.mean([x['cosine_error_no_input_error'] for x in layer_outputs_info])
            output_clip = list()
            for i, output in enumerate(layer_outputs_info):
                if output['clip_calculated']:
                    output_clip.append(output['clip_ratio'])
            output_clip = np.mean(output_clip) if len(output_clip) > 0 else -1
            if output_cos_error_acc_off > self.error_threshold.cos_sim_decrease_threshold_2 or output_cos_error_acc_off > 0.2 * output_cos_error_acc_on:
                if output_clip > self.error_threshold.clip_ratio_threshold:
                    self.low_acc_layers_error_source[idx]['Data Clip'] = 1
                if layer_type in ["conv", "convtranspose", "gemm", "fc", "depthwiseconv", "matmul"]:
                    if wegiht_noise_ratio > self.error_threshold.weight_noise_ratio_threshold or qweight_abs_mean < self.error_threshold.quan_weight_abs_mean_threshold: # type: ignore
                        self.low_acc_layers_error_source[idx]['Weight Quantization Noise'] = 1
            else:
                self.low_acc_layers_error_source[idx]['Input Accumulated Error'] = 1

    def generate_analysis_log(self):
        def add_paragraph_into_table_cell(table, i, j, contents=[]):
            if len(contents) > 1:
                for content in contents:
                    table.cell(i+1, j).add_paragraph(content)
                table.cell(i+1, j).text = table.cell(i+1, j).text.lstrip()
            elif len(contents) == 1:
                table.cell(i+1, j).text = contents[0]
            else:
                table.cell(i+1, j).text = '--'
            return table, j+1

        self.log = Document()
        # heading1
        self.log.add_heading("Model Quantization Error Analysis Info")
        # heading2
        self.log.add_heading("Layer similarity metrics", level=2)

        self.log.add_paragraph("Fig.1 L1-error \n")
        img_path = os.path.join(self.output_dir, "L1_error.png")
        image = cv2.imread(img_path)
        if image is not None:
            h, w = image.shape[:2]
            if w / h > 1.5:
                self.log.add_picture(img_path, width=Inches(6))
            else:
                self.log.add_picture(img_path, height=Inches(4))
            self.log.add_paragraph("Fig.2 L2 error )\n")

        img_path = os.path.join(self.output_dir, "L2_error.png")
        image = cv2.imread(img_path)
        if image is not None:
            h, w = image.shape[:2]
            if w / h > 1.5:
                self.log.add_picture(img_path, width=Inches(6))
            else:
                self.log.add_picture(img_path, height=Inches(4))
            self.log.add_paragraph("Fig.3 Cosine error \n")

        img_path = os.path.join(self.output_dir, "Cosine_error.png")
        image = cv2.imread(img_path)
        if image is not None:
            h, w = image.shape[:2]
            if w / h > 1.5:
                self.log.add_picture(img_path, width=Inches(6))
            else:
                self.log.add_picture(img_path, height=Inches(4))
            # heading3
            self.log.add_heading("Layer cosine similarity variance", level=2)

        self.log.add_paragraph("Fig.4 layer cosine error variance\n")
        img_path = os.path.join(self.output_dir, "Cosine_error_variance.png")
        image = cv2.imread(img_path)
        if image is not None:
            h, w = image.shape[:2]
            if w / h > 1.5:
                self.log.add_picture(img_path, width=Inches(6))
            else:
                self.log.add_picture(img_path, height=Inches(4))

        # heading 4
        self.log.add_heading("Per layer metrics", level=2)
        self.log.add_paragraph("Table 1 Per layer metrics")
        # table headers
        headers = [
            'i', 'name', 'op-type',
            'input range (gt)', 'input range (quantized)',
            'output L1 error', 'output L2 error',
            'output cosine error', 'clip range ratio',
            'txme_clip_ratio', 'txme_lr_range',
            'weight L1 error', 
        ]
        table = self.log.add_table(rows=len(self.layer_error_info) + 1,
                                   cols=len(headers),
                                   style='Table Grid')
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        # print per-layer info
        for i, info in enumerate(self.layer_error_info):
            table, j = add_paragraph_into_table_cell(table, i, j=0, contents=[str(i)])
            table, j = add_paragraph_into_table_cell(table, i, j, contents=[info['layer_name']])
            table, j = add_paragraph_into_table_cell(table, i, j, contents=[info['layer_type']])

            # inputs
            if info['layer_type'] == 'data':
                table, j = add_paragraph_into_table_cell(table, i, j, contents=['--'])
                table, j = add_paragraph_into_table_cell(table, i, j, contents=['--'])
                table, j = add_paragraph_into_table_cell(table, i, j, contents=['--'])
            else:
                keys = [['min_onnx', 'max_onnx'], ['min_quant', 'max_quant']]
                for key in keys:
                    contents = []
                    for input_info in info['inputs']:
                        content = "[%.3f, %.3f]" % (input_info[key[0]], input_info[key[1]])
                        contents.append(content)
                    table, j = add_paragraph_into_table_cell(table, i, j, contents=contents)

            # outputs
            keys = [
                'L1_error', 'L2_error', 'cosine_error'
            ]
            for key in keys:
                contents = []
                for output_info in info['outputs']:
                    content = str("%.4f" % output_info[key])
                    contents.append(content)
                table, j = add_paragraph_into_table_cell(table, i, j, contents=contents)
                
            outputs_info = info['outputs']
            clip_ratio = []
            if isinstance(outputs_info, list):
                for item in outputs_info:
                    clip_ratio.append("%.6f"%item['clip_ratio'])
            else:
                clip_ratio.append("%.6f"%outputs_info['clip_ratio'])
            table, j = add_paragraph_into_table_cell(table, i, j, contents=clip_ratio)
            
            txme_clip_ratio = []
            if isinstance(outputs_info, list):
                for item in outputs_info:
                    txme_clip_ratio.append("%.6f"%item['txme_clip_ratio'])
            else:
                txme_clip_ratio.append("%.6f"%outputs_info['txme_clip_ratio'])
            table, j = add_paragraph_into_table_cell(table, i, j, contents=txme_clip_ratio)
            
            contents = []
            for v in outputs_info[0]['txme_lr_range']:
                content = str("%d" % v)
                contents.append(content)
            table, j = add_paragraph_into_table_cell(table, i, j, contents=contents)
                                
            contents = []
            for input_info in info['inputs']:
                content = "--" if info['layer_type'] not in ["conv", "convtranspose", "gemm", "fc", "depthwiseconv", "matmul"] else "%.4f" % info['weights']['weight_L1_error']
                contents.append(content)
            table, j = add_paragraph_into_table_cell(table, i, j, contents=contents)

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(7.5)

        if len(self.low_acc_layer_idx) > 0:
            self.log.add_paragraph("The layers on which cosine similarity between the full precision's results \
                and the quantized results decrease dramatically are showing as below.\n")
            self.log.add_paragraph("Table 2. Low accuracy layers' table")
            table = self.log.add_table(rows = len(self.low_acc_layer_idx)+1, cols = 5,style='Table Grid')
            table.cell(0,0).text='layer name'
            table.cell(0,1).text='layer index'
            table.cell(0,2).text='layer type'
            table.cell(0,3).text='cosine error'
            # table.cell(0,4).text='cosine error(without input accum error)'
            table.cell(0,4).text='clip range ratio'
            for i, layer_idx in enumerate(self.low_acc_layer_idx):
                #layer_idx = low_acc_layer_idx[i]
                layer_info = self.layer_error_info[layer_idx]
                layer_name = layer_info['layer_name']
                layer_type = layer_info['layer_type']
                outputs_info = layer_info['outputs']
                cosine_err = ""
                if isinstance(outputs_info, list):
                    for item in outputs_info:
                        cosine_err += "%.6f\n"%item['cosine_error']
                else:
                    cosine_err = "%.6f\n"%outputs_info['cosine_error']
                # cosine_err_2 = layer_info['outputs'][0]['cosine_error_no_input_error']
                # info = [layer_name,layer_idx,layer_type,"%.6f"%cosine_err,"%.6f"%cosine_err_2]
                # clip_ratio = layer_info['outputs'][0]['clip_ratio']
                clip_ratio = ""
                if isinstance(outputs_info, list):
                    for item in outputs_info:
                        clip_ratio += "%.6f\n"%item['clip_ratio']
                else:
                    clip_ratio = "%.6f\n"%outputs_info['clip_ratio']
                # clip_ratio = ["%.6f"%item['clip_ratio'] for item in outputs_info] if isinstance(outputs_info, list) else "%.6f"%outputs_info['clip_ratio']
                info = [layer_name, layer_idx, layer_type, cosine_err[:-2], clip_ratio[:-2]]
                for j in range(5):
                    table.cell(i+1, j).text= str(info[j])

        self.log.add_heading("Low accuracy layer error analysis", level=2)
        if len(self.low_acc_layer_idx) > 0:
            self.log.add_paragraph("The low accuracy layers' quantization error infomation are as below.")
        else:
            self.log.add_paragraph("The network has low quantization accuracy loss, congratulations!")
        for idx, layer_idx in enumerate(self.low_acc_layer_idx):
            layer_info = self.layer_error_info[layer_idx]
            layer_name = layer_info['layer_name']
            layer_type = layer_info['layer_type']
            layer_inputs_info = layer_info['inputs']
            layer_outputs_info = layer_info['outputs']
            weight_info = None
            if layer_type in ["conv", "convtranspose", "gemm", "fc", "depthwiseconv", "matmul"]:
                weight_info = layer_info['weights']

            self.log.add_paragraph("#### layer id = %d , layer name = %s , layer type = %s"%(layer_idx, layer_name, layer_type))
            self.log.add_paragraph("---- Layer output errors: ")
            for i, output in enumerate(layer_outputs_info):
                self.log.add_paragraph("-------- Output %d : name = %s , error mean = %.6f , std = %.6f , noise ratio = %.6f , L1 error = %.6f , L2 error = %.6f , cosine dist = %.6f"\
                %(i, output['name'], output['mean_noise'], output['std_noise'], output['noise_ratio'], output["L1_error"], output["L2_error"], output["cosine_error"]))
            self.log.add_paragraph("---- Error of inputs:")
            for i, input in enumerate(layer_inputs_info):
                self.log.add_paragraph("-------- Input %d : name = %s , error mean = %.6f , std = %.6f , noise ratio = %.6f , L1 error = %.6f , L2 error = %.6f , cosine dist = %.6f"\
                %(i, input['name'], input['mean_noise'], input['std_noise'], input['noise_ratio'], input["L1_error"], input["L2_error"], input["cosine_error"]))
            self.log.add_paragraph("---- Without accumulated input error, the outputs error metrics are:")
            for i, output in enumerate(layer_outputs_info):
                self.log.add_paragraph("-------- Output %d : name = %s, L1 error = %.6f , L2 error = %.6f , cosine dist = %.6f"\
                %(i, output['name'], output["L1_error_no_input_error"], output["L2_error_no_input_error"], output["cosine_error_no_input_error"]))

            if layer_type in ["conv", "convtranspose", "gemm", "fc", "depthwiseconv", "matmul"]:
                self.log.add_paragraph("---- Error of weights:")
                self.log.add_paragraph("-------- Weight error mean = %f , std = %f , noise ratio = %.6f , quantized_weight abs-mean = %.6f"\
                    %(weight_info['noise_mean'], weight_info["noise_std"], weight_info['noise_ratio'], weight_info['qweight_abs_mean'])) # type: ignore
            self.log.add_paragraph("---- Error of data type trans clip:")
            for i, output in enumerate(layer_outputs_info):
                if output['clip_calculated']:
                    self.log.add_paragraph("---- Output %d : name = %s, clip ratio = %.6f"\
                    %(i, output['name'], output["clip_ratio"]))
                else:
                    self.log.add_paragraph("---- Output %d : name = %s, clip ratio not counted."\
                    %(i, output['name']))

            error_sources = []
            for key in self.low_acc_layers_error_source[idx].keys():
                if self.low_acc_layers_error_source[idx][key]==1:
                    error_sources.append(key)
            self.log.add_paragraph("**** Layer's error is maily due to %s:\n"%str(error_sources))
        self.log.save(self.log_file_name)
        # self.n_samples = 0

    def annlyzer(self):
        try:        
            self.create_output_folder()
            self.calc_weight_distribution()
            self.calc_featuremap_distribution()
            self.generate_layer_error_info()
            self.calc_layer_error_metrics()
            self.save_result_figures()
            self.low_accuracy_analysis()
            self.generate_analysis_log()
            self.n_samples = 0
        except:
            print("error analyzer failure!")
            os._exit(-1)

    def __call__(self, in_data, **kwargs):
        if self.n_samples == 0:
            self.weight_distribution_dict = self.calc_weight_distribution()
            self.featuremap_distribution_dict = self.init_featuremap_info_dict()
            self.layer_error_info = self.init_layer_error_list()
            self.layer_name_idx_dict = self.create_layer_name_idx_dict()
            self.acc_error = kwargs.get("acc_error", True)
            self.network_input_names = []
            for layer in self.get_layers():
                if layer.get_layer_type() == "data":
                    self.network_input_names.extend(layer.get_onnx_output_name())
        return super().__call__(in_data, **kwargs)
